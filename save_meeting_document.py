# Module để lưu nội dung cuộc họp dưới dạng PDF và WORD
# Sử dụng: python-docx cho WORD và reportlab hoặc fpdf2 cho PDF

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def sanitize_filename(name):
    """Loại bỏ các ký tự không hợp lệ trong tên file"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        name = name.replace(char, '_')
    return name


def create_meeting_folder(meeting_code):
    """Tạo thư mục cho cuộc họp nếu chưa tồn tại"""
    folder_path = os.path.join(os.path.dirname(__file__), 'meeting', sanitize_filename(meeting_code))
    os.makedirs(folder_path, exist_ok=True)
    return folder_path


def save_as_word(meeting_info, content, folder_path):
    """Lưu nội dung cuộc họp dưới dạng WORD (.docx)"""
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('BIÊN BẢN CUỘC HỌP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin cuộc họp
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    # Row 1: Tên cuộc họp
    info_table.rows[0].cells[0].text = 'Tên cuộc họp:'
    info_table.rows[0].cells[1].text = meeting_info.get('meetingName', '')
    
    # Row 2: Mã cuộc họp
    info_table.rows[1].cells[0].text = 'Mã cuộc họp:'
    info_table.rows[1].cells[1].text = meeting_info.get('meetingCode', '')
    
    # Row 3: Chủ tọa
    info_table.rows[2].cells[0].text = 'Chủ tọa:'
    info_table.rows[2].cells[1].text = meeting_info.get('hostName', '')
    
    # Row 4: Thư ký
    info_table.rows[3].cells[0].text = 'Thư ký:'
    info_table.rows[3].cells[1].text = meeting_info.get('secretaryName', '')
    
    # Bold các label
    for row in info_table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    # Thời gian tạo
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_para.add_run(f"Thời gian: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    # Nội dung cuộc họp
    doc.add_paragraph()
    doc.add_heading('NỘI DUNG CUỘC HỌP', level=1)
    doc.add_paragraph()
    
    # Xử lý nội dung - loại bỏ các dòng hệ thống
    lines = content.split('\n')
    cleaned_lines = []
    for line in lines:
        # Bỏ qua các dòng thông báo hệ thống
        if line.startswith('[') and line.endswith(']'):
            continue
        if line.strip():
            cleaned_lines.append(line)
    
    content_para = doc.add_paragraph('\n'.join(cleaned_lines))
    content_para.paragraph_format.line_spacing = 1.5
    
    # Lưu file
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"bien_ban_{timestamp}.docx"
    filepath = os.path.join(folder_path, filename)
    doc.save(filepath)
    
    return filepath


class MeetingPDF(FPDF):
    """Custom PDF class với hỗ trợ Unicode tiếng Việt"""
    
    def header(self):
        pass
    
    def footer(self):
        self.set_y(-15)
        self.set_font('DejaVu', '', 8)
        self.cell(0, 10, f'Trang {self.page_no()}', 0, 0, 'C')


def save_as_pdf(meeting_info, content, folder_path):
    """Lưu nội dung cuộc họp dưới dạng PDF"""
    pdf = MeetingPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Thêm font hỗ trợ tiếng Việt
    font_dir = os.path.join(os.path.dirname(__file__), 'fonts')
    
    # Kiểm tra và sử dụng font DejaVu nếu có, nếu không dùng font mặc định
    font_path = os.path.join(font_dir, 'DejaVuSans.ttf')
    font_bold_path = os.path.join(font_dir, 'DejaVuSans-Bold.ttf')
    
    if os.path.exists(font_path):
        pdf.add_font('DejaVu', '', font_path, uni=True)
        pdf.add_font('DejaVu', 'B', font_bold_path, uni=True)
        use_unicode = True
    else:
        # Fallback - tạo thư mục fonts và thông báo
        os.makedirs(font_dir, exist_ok=True)
        use_unicode = False
        print(f"[WARNING] Để hỗ trợ tiếng Việt trong PDF, hãy tải font DejaVuSans.ttf vào thư mục {font_dir}")
    
    pdf.add_page()
    
    if use_unicode:
        # Tiêu đề
        pdf.set_font('DejaVu', 'B', 18)
        pdf.cell(0, 15, 'BIÊN BẢN CUỘC HỌP', 0, 1, 'C')
        pdf.ln(5)
        
        # Thông tin cuộc họp
        pdf.set_font('DejaVu', '', 12)
        pdf.cell(50, 10, 'Tên cuộc họp:', 0)
        pdf.cell(0, 10, meeting_info.get('meetingName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Mã cuộc họp:', 0)
        pdf.cell(0, 10, meeting_info.get('meetingCode', ''), 0, 1)
        
        pdf.cell(50, 10, 'Chủ tọa:', 0)
        pdf.cell(0, 10, meeting_info.get('hostName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Thư ký:', 0)
        pdf.cell(0, 10, meeting_info.get('secretaryName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Thời gian:', 0)
        pdf.cell(0, 10, datetime.now().strftime('%d/%m/%Y %H:%M:%S'), 0, 1)
        
        pdf.ln(10)
        
        # Nội dung cuộc họp
        pdf.set_font('DejaVu', 'B', 14)
        pdf.cell(0, 10, 'NỘI DUNG CUỘC HỌP', 0, 1)
        pdf.ln(5)
        
        pdf.set_font('DejaVu', '', 11)
    else:
        # Fallback với font cơ bản (không hỗ trợ Unicode tốt)
        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 15, 'BIEN BAN CUOC HOP', 0, 1, 'C')
        pdf.ln(5)
        
        pdf.set_font('Arial', '', 12)
        pdf.cell(50, 10, 'Ten cuoc hop:', 0)
        pdf.cell(0, 10, meeting_info.get('meetingName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Ma cuoc hop:', 0)
        pdf.cell(0, 10, meeting_info.get('meetingCode', ''), 0, 1)
        
        pdf.cell(50, 10, 'Chu toa:', 0)
        pdf.cell(0, 10, meeting_info.get('hostName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Thu ky:', 0)
        pdf.cell(0, 10, meeting_info.get('secretaryName', ''), 0, 1)
        
        pdf.cell(50, 10, 'Thoi gian:', 0)
        pdf.cell(0, 10, datetime.now().strftime('%d/%m/%Y %H:%M:%S'), 0, 1)
        
        pdf.ln(10)
        
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'NOI DUNG CUOC HOP', 0, 1)
        pdf.ln(5)
        
        pdf.set_font('Arial', '', 11)
    
    # Xử lý nội dung - loại bỏ các dòng hệ thống
    lines = content.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.startswith('[') and line.endswith(']'):
            continue
        if line.strip():
            cleaned_lines.append(line)
    
    content_text = '\n'.join(cleaned_lines)
    pdf.multi_cell(0, 7, content_text)
    
    # Lưu file
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"bien_ban_{timestamp}.pdf"
    filepath = os.path.join(folder_path, filename)
    pdf.output(filepath)
    
    return filepath


def save_meeting_documents(meeting_info, content):
    """
    Lưu nội dung cuộc họp dưới cả 2 định dạng PDF và WORD
    
    Args:
        meeting_info: dict chứa thông tin cuộc họp (meetingName, meetingCode, hostName, secretaryName)
        content: string nội dung transcript của cuộc họp
    
    Returns:
        dict với đường dẫn các file đã lưu
    """
    meeting_code = meeting_info.get('meetingCode', 'unknown')
    folder_path = create_meeting_folder(meeting_code)
    
    result = {
        'folder': folder_path,
        'word': None,
        'pdf': None,
        'errors': []
    }
    
    # Lưu file WORD
    try:
        result['word'] = save_as_word(meeting_info, content, folder_path)
        print(f"[OK] Đã lưu file WORD: {result['word']}")
    except Exception as e:
        error_msg = f"Lỗi khi lưu file WORD: {str(e)}"
        result['errors'].append(error_msg)
        print(f"[ERROR] {error_msg}")
    
    # Lưu file PDF
    try:
        result['pdf'] = save_as_pdf(meeting_info, content, folder_path)
        print(f"[OK] Đã lưu file PDF: {result['pdf']}")
    except Exception as e:
        error_msg = f"Lỗi khi lưu file PDF: {str(e)}"
        result['errors'].append(error_msg)
        print(f"[ERROR] {error_msg}")
    
    return result


if __name__ == "__main__":
    # Test
    test_info = {
        'meetingName': 'Cuộc họp test',
        'meetingCode': 'TEST001',
        'hostName': 'Nguyễn Văn A',
        'secretaryName': 'Trần Thị B'
    }
    test_content = """[Cuộc họp đã bắt đầu]
[Mic đã bật]
Xin chào các bạn, hôm nay chúng ta sẽ thảo luận về dự án mới.
Đầu tiên, tôi muốn giới thiệu về mục tiêu của dự án.
Chúng ta cần hoàn thành trong 3 tháng tới.
[Mic đã tắt]
[Cuộc họp đã kết thúc]"""
    
    result = save_meeting_documents(test_info, test_content)
    print(f"\nKết quả: {result}")
